# report_exporter.py
from flask import send_file, request
from io import BytesIO
from fpdf import FPDF
import pandas as pd
from openpyxl import Workbook
from docx import Document
from app import app, es

def fetch_logs():
    honeypot = request.args.get("honeypot", "*")
    start_date = request.args.get("start_date")
    end_date = request.args.get("end_date")
    query = {"size": 500, "query": {"bool": {"must": []}}}
    if honeypot != "*":
        query["query"]["bool"]["must"].append({"match": {"program": honeypot}})
    if start_date and end_date:
        query["query"]["bool"]["must"].append({
            "range": {"@timestamp": {"gte": start_date, "lte": end_date}}
        })
    results = es.search(index="logstash-*", body=query)
    return [hit["_source"] for hit in results["hits"]["hits"]]

@app.route("/api/export/pdf")
def export_pdf():
    logs = fetch_logs()
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=10)
    pdf.cell(200, 10, txt="Honeypot Logs", ln=True, align="C")
    for log in logs:
        msg = f"{log.get('@timestamp')} | {log.get('src_ip')} | {log.get('proto')} | {log.get('message')}"
        pdf.multi_cell(0, 10, txt=msg)
    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="report.pdf")

@app.route("/api/export/csv")
def export_csv():
    logs = fetch_logs()
    df = pd.DataFrame(logs)
    buffer = BytesIO()
    df.to_csv(buffer, index=False)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="report.csv", mimetype='text/csv')

@app.route("/api/export/excel")
def export_excel():
    logs = fetch_logs()
    df = pd.DataFrame(logs)
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="report.xlsx")

@app.route("/api/export/word")
def export_word():
    logs = fetch_logs()
    doc = Document()
    doc.add_heading("Honeypot Logs Report", level=1)
    for log in logs:
        entry = f"{log.get('@timestamp')} - {log.get('src_ip')} - {log.get('proto')} - {log.get('message')}"
        doc.add_paragraph(entry)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="report.docx")
